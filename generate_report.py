from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
import os

# 创建文档
doc = Document()

def add_section_heading(doc, text, level=1):
    h = doc.add_heading(level=level)
    run = h.add_run(text)
    run.font.size = Pt(16 if level == 1 else 14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)
    return h

def add_body(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.bold = bold
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_after = Pt(6)
    return p

def add_image(doc, img_path, width=Inches(5.5)):
    if os.path.exists(img_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(img_path, width=width)
        p.paragraph_format.space_after = Pt(12)
        return True
    print(f"Image not found: {img_path}")
    return False

# 标题
title = doc.add_heading("尤瑞克公司及其南非拉霍马矿山项目综合报告", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 128)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("基于《ULRIC 2026 PPT-双语版》与《南非拉霍马矿山商业计划书V3》")
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph()

# 一、公司概况
add_section_heading(doc, "一、公司概况")
add_body(doc, "尤瑞克进出口有限公司（Ulric Import Export Limited）于2016年3月1日在加拿大阿尔伯塔省卡尔加里注册成立，商业注册号2019540141。经过多年发展，公司已逐步成长为集矿产投资、商品贸易、家具研发设计为一体的多元化现代企业，在加拿大及海外拥有众多战略合作伙伴，并持有Amor Vincit和Philly两个国际品牌商标。")
add_body(doc, "公司核心业务涵盖四大板块：南非矿山开发、南非矿石贸易、美式家具研发与设计、北美木材贸易。其中，南非矿山与矿石贸易是公司的战略性重点业务。")

add_image(doc, r"C:\Users\lk\pdf_images\preview_ppt_p3.png")
add_body(doc, "图1：尤瑞克公司介绍", bold=True)

add_body(doc, "在矿石贸易方面，公司凭借南非独特的矿产资源优势，通过经验丰富的贸易团队、先进的技术、可靠的服务和成熟的销售渠道，服务于全球钢铁冶金和钒储能行业。主要产品包括南非钒钛磁铁矿、赤铁矿、铬矿、铜矿、锰矿等。")

add_image(doc, r"C:\Users\lk\pdf_images\preview_ppt_p8.png")
add_body(doc, "图2：公司四大主营业务", bold=True)

# 二、拉霍马矿山项目核心信息
add_section_heading(doc, "二、拉霍马矿山项目核心信息")
add_body(doc, "拉霍马矿业资源有限公司（Rakhoma Mining Resources (Pty) Ltd）持有的矿山位于南非林波波省色库昆区伯格斯福特镇西南约55公里处，斯迪尔伯特镇以西20公里。矿山采矿权区占地3165.32公顷，包括格鲁克512KS矿场、格鲁克欧斯513KS矿场和铁矿石847KS矿场，开采区域集中在约395公顷范围内。")

add_image(doc, r"C:\Users\lk\pdf_images\preview_bp_p4.png")
add_body(doc, "图3：矿山项目区域位置平面图", bold=True)

add_body(doc, "矿区地质条件优越，矿层产状以水平产状为主，存在两组主要断裂带。主矿层上覆辉长岩-苏长岩，矿体直接出露地表，属露天开采，倾角在4度–10度之间，最大开采深度约20米。")

add_image(doc, r"C:\Users\lk\pdf_images\preview_bp_p5.png")
add_body(doc, "图4：矿区地质与露头矿层", bold=True)

add_body(doc, "矿山拥有约6000万吨储量的钒钛磁铁矿资源，主矿层平均厚度达3.19米，钒品位最高可达1.7%以上。经磁选后钒品位可提升至1.64%。地表浅部矿石铁品位在54%以上，钒含量在1.5%以上。高品位矿石合计约1.3亿吨，剥采比仅0.07，开采条件极佳。扣除地质损失后的资源评估显示，主矿层实测储量超1100万吨，指明储量近2000万吨。")

add_image(doc, r"C:\Users\lk\pdf_images\preview_bp_p6.png")
add_body(doc, "图5：岩心钻探钻孔取芯图", bold=True)

# 三、采选工程与运营规划
add_section_heading(doc, "三、采选工程与运营规划")
add_body(doc, "项目规划矿山采选规模为年处理原矿360万吨，年产钒钛磁铁矿石产品300万吨，矿山预计寿命20年。产品指标为：TFe 53–55%、V2O5 1.6–1.7%、TiO2 11–13%。")
add_body(doc, "矿山采用露天山坡开采方式，沿走向开采约20米宽的条带。采矿方法包括条带开采和同步复垦，工艺涵盖穿孔、爆破、采装、运输等工序。在有植被地区，表层土壤会被堆存以便后期复垦使用。")

add_image(doc, r"C:\Users\lk\pdf_images\preview_bp_p10.png")
add_body(doc, "图6：采矿方法与工程规划", bold=True)

add_body(doc, "选矿工艺采用成熟的破碎—筛分—磁选流程：原矿经格筛、颚式破碎机粗碎、圆锥破碎机细碎、双层振动筛分后，+8–40mm产品作为最终块矿，-8mm矿石经磁选后获得粉矿精矿。")

add_image(doc, r"C:\Users\lk\pdf_images\preview_bp_p12.png")
add_body(doc, "图7：选矿工艺示意图", bold=True)

add_body(doc, "物流方面，产品由自卸卡车从矿区运输至罗斯内卡尔铁路支线产品储存区，经陆运至理查德湾港（约640公里）或莫桑比克马普托港（约420公里），再海运至中国各大港口。")

# 四、经济效益分析
add_section_heading(doc, "四、经济效益分析")
add_body(doc, "按平均到岸价格100美元/吨计算，300万吨产品年销售总额达3亿美元。矿石生产成本约为26美元/吨产品，中国到岸成本约95美元/吨，吨产品利润约5美元，年利润总额1500万美元。项目投资回收期仅2年，抗风险能力强。")

add_image(doc, r"C:\Users\lk\pdf_images\preview_bp_p16.png")
add_body(doc, "图8：投资及效益分析", bold=True)

# 五、市场分析
add_section_heading(doc, "五、市场分析")
add_body(doc, "成品钒钛磁铁矿主要销往中国各大港口，用于湿法提钒及火法钢铁生产企业。目标客户包括中国大型钢铁企业（如建龙钢铁、攀枝花钢铁）及南非本地钢铁、钒产品生产企业（如VVP/WANCHEM等）。")
add_body(doc, "含钒钢具有强度高、韧性大、耐磨性好等优良特性，钢铁行业钒用量约占钒消耗总量的90%。与此同时，全钒液流电池储能技术正成为规模储能的首选技术之一，具有容量大、寿命长、瞬间充电、安全可靠、环保可回收等优势。")
add_body(doc, "市场前景极为广阔：预计2026年国内全钒液流电池新增装机规模将达2.4GW；2030年国内累计装机规模将达24GW；2031年全球钒液流电池装机容量将达32.8GWh。按1KWh电解液需8–12kg五氧化二钒计算，2030年国内高纯钒需求将达24万吨以上。")

add_image(doc, r"C:\Users\lk\pdf_images\preview_bp_p24.png")
add_body(doc, "图9：钒电池市场前景", bold=True)

# 六、产业链优势
add_section_heading(doc, "六、产业链优势")
add_body(doc, '尤瑞克已构建起从矿山到终端产品的完整钒产业链。上游涵盖矿山收购并购、合作委托、勘探开采；中游包括内陆物流（公路/铁路联运）、海运、到岸、报关清关、仓储物流、港口销售；下游深加工分为火法路线（高炉/电炉→钒渣→五氧化二钒→含钒合金）和湿法路线（回转窑湿法提钒→偏钒酸铵→五氧化二钒→钒储能电解液）。')
add_body(doc, '公司具备专业涉外物流团队、贸易融资团队和稳定的下游终端合作用户，形成长期稳定的"包采包销"模式。')

add_image(doc, r"C:\Users\lk\pdf_images\preview_ppt_p18.png")
add_body(doc, "图10：完善的钒产业链及专业操作流程", bold=True)

# 七、项目亮点与合作愿景
add_section_heading(doc, "七、项目亮点与合作愿景")
add_body(doc, "拉霍马矿山项目具备六大核心亮点：")
add_body(doc, "A. 矿山所在地区工业基础设施雄厚，水、电、路及人力资源齐备，具备随时开发条件；")
add_body(doc, "B. 采矿权证等资质许可齐全，项目方拥有完整的矿山开采权及完备的产业相关手续；")
add_body(doc, "C. 矿体直接出露地表，埋藏浅且矿体厚大，采用露天开采方式，剥采比低，属易选矿石；")
add_body(doc, "D. 矿山潜在资源量巨大，推测储量远超亿吨；")
add_body(doc, "E. 矿石产品属于低硫、低磷、高钒的优质钒钛磁铁矿；")
add_body(doc, "F. 项目方经多年深耕，收获了良好的公共关系和社区关系，夯实了矿山稳定开发的长期经营基础。")

add_image(doc, r"C:\Users\lk\pdf_images\preview_bp_p26.png")
add_body(doc, "图11：项目亮点", bold=True)

add_body(doc, "公司愿景是依托钒钛磁铁矿资源优势，发挥专家团队优势，利用渠道、资金和服务优势，打造独具特色的、年产值过百亿的储能新材料供应链，为中国新能源的利用与发展做贡献，并建设具有核心竞争力的火法/湿法冶炼生产企业。")

add_image(doc, r"C:\Users\lk\pdf_images\preview_bp_p28.png")
add_body(doc, "图12：愿景规划与合作邀请", bold=True)

add_image(doc, r"C:\Users\lk\pdf_images\preview_ppt_p25.png")
add_body(doc, "图13：尤瑞克公司使命、愿景与价值观", bold=True)

# 八、总结评价
add_section_heading(doc, "八、总结评价")
add_body(doc, '综合来看，尤瑞克公司及其拉霍马矿山项目展现出"资源禀赋优越 + 产业链完整 + 市场前景广阔"的综合竞争力：')
add_body(doc, "1. 资源端：南非拉霍马矿山拥有世界级高品位钒钛磁铁矿资源，露天开采条件极佳，剥采比仅0.07，生产成本极具竞争力，且潜在推测储量远超亿吨。")
add_body(doc, "2. 产业端：公司已构建从矿山开采、物流运输到深加工的完整产业链，尤其在钒储能领域具备前瞻性布局，同时拥有专业的涉外物流、贸易融资和业务团队。")
add_body(doc, "3. 市场端：传统钢铁市场提供稳定基本盘，而钒液流电池储能市场的爆发式增长（预计2030年全球累计装机超30GWh）将带来巨大的增量需求，市场空间广阔。")
add_body(doc, "4. 财务端：项目年处理原矿360万吨，年产成品300万吨，年销售总额3亿美元，年利润1500万美元，投资回收期仅2年，经济效益显著，抗风险能力强。")
add_body(doc, "整体而言，该项目是一个资源稀缺性强、产业链整合度高、契合全球新能源转型趋势的优质矿产投资项目，具备良好的投资价值与合作前景。")

# 保存文档
output_path = r"C:\Users\lk\Desktop\尤瑞克公司及拉霍马矿山项目综合报告.docx"
doc.save(output_path)
print(f"Document saved to: {output_path}")
